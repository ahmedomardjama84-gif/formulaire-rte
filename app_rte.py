"""
Formulaire collaboratif - Grille de maturite RTE
Application Streamlit avec notation automatique :
  - Chaque question : Observe (1pt) / Partiel (0.5pt) / Non observe (0pt)
  - Score critere = (somme points / nb questions) * 4
  - Score axe = moyenne des criteres
  - Score global = moyenne des criteres

Lancement local :
    streamlit run app_rte.py
"""

import sqlite3
import json
from datetime import datetime, date
from io import BytesIO
import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIGURATION
# ============================================================

DB_PATH = "reponses_rte.db"
APP_TITLE = "Grille de maturite RTE - Formulaire collaboratif"
APP_ICON = "📊"

# Code secret admin : modifiable ici OU via Streamlit secrets (cle "admin_code")
# Les collegues voient seulement le formulaire ; toi tu accedes a tout avec ?admin=<code>
DEFAULT_ADMIN_CODE = "rte2026"

def get_admin_code():
    """Renvoie le code admin (priorite a st.secrets, sinon fallback)."""
    try:
        return st.secrets.get("admin_code", DEFAULT_ADMIN_CODE)
    except Exception:
        return DEFAULT_ADMIN_CODE

AXE_COLORS = {
    1: "#C28533",  # orange
    2: "#1F4E79",  # bleu
    3: "#256D5A",  # vert
    4: "#5A3B7A",  # violet
    5: "#9C2A2A",  # rouge
}

POINTS = {"obs": 1.0, "part": 0.5, "non": 0.0}
STATUT_LABELS = {
    "obs": "Observe",
    "part": "Partiel",
    "non": "Non observe",
}
STATUT_LABELS_VISUAL = {
    "obs": "✅ Observe",
    "part": "🟡 Partiel",
    "non": "❌ Non observe",
}
STATUT_COLORS = {
    "obs": "#2E8B73",
    "part": "#C9A961",
    "non": "#C0392B",
}

AXES = {
    1: "Diagnostic territorial",
    2: "Vision territoriale et recits",
    3: "Cooperation territoriale",
    4: "Gouvernance inclusive",
    5: "Redistribuer et transformer",
}

CRITERES = [
    ("1.1", 1, "Specification et diagnostic du territoire"),
    ("1.2", 1, "Modelisation des flux territoriaux"),
    ("1.3", 1, "Redevabilites : double materialite (impact territoire <-> organisation)"),
    ("2.1", 2, "Vision territoriale d'ensemble et vision sectorielle"),
    ("2.2", 2, "Besoins des parties prenantes internes et des usagers"),
    ("2.3", 2, "Besoins des communautes peripheriques"),
    ("2.4", 2, "Integration du vivant et des ecosystemes"),
    ("2.5", 2, "Recits de territoire et identite culturelle"),
    ("3.1", 3, "Culture et competences de cooperation territoriale"),
    ("3.2", 3, "Demarches collectives multi-echelles et partenariats formalises"),
    ("4.1", 4, "Gouvernance interne inclusive et transparente"),
    ("4.2", 4, "Processus de concertation et de dialogue"),
    ("4.3", 4, "Evaluation collective et redevabilite partagee"),
    ("5.1", 5, "Reorientation des echanges vers l'economie locale"),
    ("5.2", 5, "Partage de la valeur produite"),
    ("5.3", 5, "Inclusion des communautes peripheriques"),
    ("5.4", 5, "Bilan des transformations : contribution au bien commun"),
]

# Checklists : la base. Chaque item devient une question notee.
CHECKLISTS_RAW = {
    "1.1": [
        "Un diagnostic territorial formalise est-il disponible (ecosysteme, etendue, diversite) ?",
        "Le diagnostic integre-t-il des sources officielles (INSEE, DRAAF, ARS, ORS, etc.) ?",
        "Le diagnostic est-il mis a jour regulierement (au moins tous les 3 ans) ?",
        "Le diagnostic est-il partage avec les parties prenantes du territoire ?",
    ],
    "1.2": [
        "Une cartographie des flux entrants/sortants (economiques, humains, materiels) existe-t-elle ?",
        "L'organisation sait-elle quels metiers et emplois font vivre le territoire ?",
        "L'organisation a-t-elle etabli une cartographie des acteurs locaux ?",
        "L'organisation suit-elle dans le temps la part de son activite qui est locale ?",
    ],
    "1.3": [
        "L'impact du territoire sur l'organisation est-il documente (ressources, contraintes, opportunites) ?",
        "L'impact de l'organisation sur le territoire est-il documente ?",
        "Un bilan de ces redevabilites bidirectionnelles est-il produit ?",
        "Ces redevabilites sont-elles discutees en gouvernance collective ?",
    ],
    "2.1": [
        "Une vision territoriale d'ensemble est-elle referencee par l'organisation ?",
        "Une vision sectorielle propre a l'organisation est-elle formalisee ?",
        "L'articulation entre vision d'ensemble et vision sectorielle est-elle explicite ?",
        "La vision est-elle co-construite avec les parties prenantes ?",
        "La vision integre-t-elle les transitions ecologiques, sociales, economiques et politiques ?",
    ],
    "2.2": [
        "Les besoins des salaries sont-ils identifies et documentes ?",
        "Les besoins des usagers (beneficiaires directs) sont-ils identifies et documentes ?",
        "Des dispositifs de recueil des besoins existent-ils (enquetes, entretiens, ateliers) ?",
        "Les besoins identifies influencent-ils les orientations strategiques ?",
    ],
    "2.3": [
        "Les populations saisonnieres sont-elles prises en compte dans la strategie ?",
        "Les publics fragiles beneficient-ils de dispositifs dedies ?",
        "Les jeunes sont-ils associes aux demarches de l'organisation ?",
        "Les acteurs eloignes ou peu visibles (ESS, artisans, collectifs) sont-ils identifies et associes ?",
    ],
    "2.4": [
        "Les milieux naturels et la biodiversite sont-ils identifies comme enjeux strategiques ?",
        "Un tableau de bord environnemental existe-t-il ?",
        "Des actions d'attenuation des changements climatiques sont-elles engagees ?",
        "Des actions d'adaptation aux changements climatiques sont-elles engagees ?",
        "Des partenariats existent-ils avec des acteurs de l'environnement (parcs, ONF, LPO, etc.) ?",
    ],
    "2.5": [
        "Le projet mobilise-t-il explicitement des recits de territoire ?",
        "L'identite culturelle locale (patrimoine materiel et immateriel) est-elle integree ?",
        "Les recits sont-ils co-construits avec les acteurs locaux ?",
        "Les recits incluent-ils la diversite des acteurs ?",
    ],
    "3.1": [
        "L'organisation a-t-elle une culture explicite de cooperation ?",
        "Les salaries sont-ils formes aux pratiques cooperatives ?",
        "L'organisation promeut-elle une communication horizontale et transversale ?",
        "Les pratiques d'inter-organisation (ateliers communs, temps partages) sont-elles regulieres ?",
    ],
    "3.2": [
        "L'organisation participe-t-elle a au moins un autre dispositif local ?",
        "L'organisation participe-t-elle a au moins 1 dispositif a l'echelle departementale ?",
        "L'organisation participe-t-elle a au moins 1 dispositif a l'echelle regionale ?",
        "L'organisation est-elle impliquee dans des demarches nationales ou europeennes ?",
        "Le role de l'organisation dans ces dispositifs est-il documente ?",
        "L'organisation a-t-elle formalise par convention ses partenariats ?",
    ],
    "4.1": [
        "La gouvernance inclut-elle au moins 3 categories de parties prenantes differentes ?",
        "Les parties prenantes representees couvrent-elles l'ensemble du territoire d'action ?",
        "Les regles de decision sont-elles documentees et transparentes ?",
        "Les salaries sont-ils representes dans les instances decisionnelles ?",
    ],
    "4.2": [
        "Des espaces de concertation formalises existent-ils ?",
        "La frequence de concertation est-elle reguliere (au moins 2 fois par an) ?",
        "Les concertations sont-elles ouvertes aux acteurs au-dela des seuls adherents ?",
        "Une restitution des concertations est-elle produite et diffusee ?",
    ],
    "4.3": [
        "L'organisation produit-elle un bilan territorial qui rend compte au territoire de son impact ?",
        "Le bilan mobilise-t-il des indicateurs verifiables ?",
        "Le bilan territorial est-il co-construit avec les parties prenantes ?",
        "Le bilan territorial est-il rendu public et debattu avec les acteurs du territoire ?",
    ],
    "5.1": [
        "Les fournisseurs locaux sont-ils prioritaires dans les achats ?",
        "Les acteurs de l'ESS sont-ils prioritaires dans les echanges ?",
        "La question de l'emploi sur la filiere locale est-elle documentee ?",
        "Un suivi chiffre des achats locaux est-il produit annuellement ?",
    ],
    "5.2": [
        "La valeur produite est-elle explicitee ?",
        "Le partage de la valeur entre les parties prenantes a-t-il ete discute ?",
        "L'organisation contribue-t-elle au financement de programmes partages sur le territoire ?",
        "Un dispositif philanthropique (mecenat, dons, fondation) est-il engage ?",
        "Les benefices ou excedents sont-ils redistribues vers le territoire ?",
    ],
    "5.3": [
        "Des actions d'inclusion economique ciblent-elles les communautes peripheriques ?",
        "Des actions d'inclusion sociale ciblent-elles ces communautes ?",
        "Les communautes peripheriques sont-elles associees a la conception des actions qui les concernent ?",
        "Des partenariats operationnels existent-ils avec France Travail / France Services / Missions Locales ?",
    ],
    "5.4": [
        "Transformation economique observable (emplois, valeur retenue localement, nouvelles filieres) ?",
        "Transformation politique observable (nouveaux espaces de dialogue, participation accrue, influence) ?",
        "Transformation sociale observable (lien social, acces aux services, inclusion mesurable) ?",
        "Transformation ecologique observable (empreinte reduite, ecosystemes, adaptation climatique) ?",
    ],
}

# Genere les checklists avec IDs (1.1.1, 1.1.2, ...)
CHECKLISTS = {
    c_id: [
        {"id": f"{c_id}.{i+1}", "label": label}
        for i, label in enumerate(items)
    ]
    for c_id, items in CHECKLISTS_RAW.items()
}

NB_CRITERES_PAR_AXE = {1: 3, 2: 5, 3: 2, 4: 3, 5: 4}
TOTAL_CRITERES = sum(NB_CRITERES_PAR_AXE.values())  # 17


# ============================================================
# BASE DE DONNEES SQLITE
# ============================================================

def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute("PRAGMA journal_mode=WAL")
    return conn

def init_db():
    conn = get_conn()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS reponses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date_soumission TEXT NOT NULL,
            nom_repondant TEXT NOT NULL,
            fonction TEXT,
            organisation TEXT NOT NULL,
            territoire TEXT,
            email TEXT,
            commentaire_global TEXT,
            donnees_json TEXT NOT NULL,
            schema_version INTEGER DEFAULT 2
        )
    """)
    # Migration : ajout colonne schema_version si manquante
    cur = conn.execute("PRAGMA table_info(reponses)")
    cols = [r[1] for r in cur.fetchall()]
    if "schema_version" not in cols:
        conn.execute("ALTER TABLE reponses ADD COLUMN schema_version INTEGER DEFAULT 1")
    conn.commit()
    conn.close()

def save_response(meta, criteres_data):
    conn = get_conn()
    conn.execute("""
        INSERT INTO reponses
        (date_soumission, nom_repondant, fonction, organisation, territoire,
         email, commentaire_global, donnees_json, schema_version)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        datetime.now().isoformat(timespec="seconds"),
        meta["nom"],
        meta.get("fonction", ""),
        meta["organisation"],
        meta.get("territoire", ""),
        meta.get("email", ""),
        meta.get("commentaire_global", ""),
        json.dumps(criteres_data, ensure_ascii=False),
        2,
    ))
    conn.commit()
    conn.close()

def load_responses():
    conn = get_conn()
    df = pd.read_sql_query("SELECT * FROM reponses ORDER BY date_soumission DESC", conn)
    conn.close()
    # Ne conserver que les reponses au schema V2 (nouveau format)
    if "schema_version" in df.columns:
        df = df[df["schema_version"] == 2].reset_index(drop=True)
    return df

def delete_response(rid):
    conn = get_conn()
    conn.execute("DELETE FROM reponses WHERE id = ?", (rid,))
    conn.commit()
    conn.close()

def reset_db():
    conn = get_conn()
    conn.execute("DELETE FROM reponses")
    conn.commit()
    conn.close()


# ============================================================
# CALCUL DES SCORES (notation automatique)
# ============================================================

def score_critere(questions_data):
    """
    Score automatique d'un critere sur /4.
    questions_data : {q_id: {'statut': 'obs'|'part'|'non', 'preuve': str}}
    """
    pts = []
    for q_id, q_data in (questions_data or {}).items():
        s = q_data.get("statut")
        if s in POINTS:
            pts.append(POINTS[s])
    if not pts:
        return None  # aucune question repondue
    return round((sum(pts) / len(pts)) * 4, 2)

def nb_questions_repondues(questions_data):
    return sum(1 for q in (questions_data or {}).values()
               if q.get("statut") in POINTS)

def score_axe(criteres_data, axe):
    """Moyenne des scores des criteres de cet axe (criteres non repondus ignores)."""
    scores = []
    for c_id, c_axe, _ in CRITERES:
        if c_axe != axe:
            continue
        c_data = criteres_data.get(c_id, {})
        questions = c_data.get("questions", {})
        s = score_critere(questions)
        if s is not None:
            scores.append(s)
    return round(sum(scores) / len(scores), 2) if scores else None

def score_global(criteres_data):
    """Moyenne des scores de tous les criteres (les non repondus sont ignores)."""
    scores = []
    for c_id, _, _ in CRITERES:
        c_data = criteres_data.get(c_id, {})
        questions = c_data.get("questions", {})
        s = score_critere(questions)
        if s is not None:
            scores.append(s)
    return round(sum(scores) / len(scores), 2) if scores else None

def qualif(score):
    if score is None: return "Non evalue"
    if score < 1.5: return "1 - Tres faible / inexistant"
    if score < 2.5: return "2 - Faible / ponctuel"
    if score < 3.5: return "3 - Moyen / structure"
    return "4 - Fort / strategique"


# ============================================================
# AGREGATION DES REPONSES
# ============================================================

def parse_donnees(donnees_json):
    return json.loads(donnees_json)

def aggregate_responses(df):
    """
    Retourne pour chaque critere :
        - score_moyen (moyenne des scores individuels)
        - score_min, score_max
        - nb_reponses (nb de personnes ayant repondu au critere)
        - questions : pour chaque question, distribution des statuts + preuves collectees
    """
    result = {}
    for c_id, _, _ in CRITERES:
        # Scores individuels sur ce critere
        scores_indiv = []
        # Pour chaque question : compteurs + preuves
        questions_agg = {
            q["id"]: {"obs": 0, "part": 0, "non": 0, "preuves": []}
            for q in CHECKLISTS[c_id]
        }
        for _, row in df.iterrows():
            d = parse_donnees(row["donnees_json"])
            c_data = d.get(c_id, {})
            questions = c_data.get("questions", {})
            s = score_critere(questions)
            if s is not None:
                scores_indiv.append(s)
            for q_id, q_data in questions.items():
                if q_id not in questions_agg:
                    continue
                statut = q_data.get("statut")
                if statut in POINTS:
                    questions_agg[q_id][statut] += 1
                preuve = (q_data.get("preuve") or "").strip()
                if preuve:
                    questions_agg[q_id]["preuves"].append(
                        f"[{row['nom_repondant']}] {preuve}"
                    )
        result[c_id] = {
            "score_moyen": round(sum(scores_indiv)/len(scores_indiv), 2) if scores_indiv else None,
            "score_min": min(scores_indiv) if scores_indiv else None,
            "score_max": max(scores_indiv) if scores_indiv else None,
            "nb_reponses": len(scores_indiv),
            "questions": questions_agg,
        }
    return result


# ============================================================
# EXPORT WORD
# ============================================================

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)

def add_para_styled(container, text, size=10, bold=False, italic=False,
                    color_hex=None, align=None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    return p

def set_cell_border(cell, color_hex="BFBFBF", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _commentaire_axe(axe_num, score, scores_moyens, agg, df):
    """Genere un commentaire synthetique auto pour un axe."""
    if score is None:
        return "Non evalue par les repondants."

    # Trouve les criteres faibles et forts de cet axe
    criteres_axe = [(c, t) for c, a, t in CRITERES if a == axe_num]
    crit_scores = []
    for c_id, titre in criteres_axe:
        s = agg[c_id]["score_moyen"]
        if s is not None:
            crit_scores.append((c_id, titre, s))

    # Compare aux autres axes
    valid_axes = {a: s for a, s in scores_moyens.items() if s is not None}
    rang = sorted(valid_axes.values(), reverse=True).index(score) + 1 if score in valid_axes.values() else None

    parts = []

    # Verdict global
    if score >= 3.5:
        parts.append("Point fort majeur du profil RTE")
    elif score >= 2.5:
        parts.append("Axe a un niveau de structuration moyen")
    elif score >= 1.5:
        parts.append("Axe ponctuel a renforcer")
    else:
        parts.append("Axe tres faible, demarche a engager")

    # Si plusieurs axes : indique le rang
    if rang and len(valid_axes) >= 3:
        if rang == 1:
            parts.append("(score le plus eleve)")
        elif rang == len(valid_axes):
            parts.append("(score le plus faible)")

    # Critere a renforcer
    if crit_scores:
        crit_faible = min(crit_scores, key=lambda x: x[2])
        crit_fort = max(crit_scores, key=lambda x: x[2])
        if crit_faible[2] < crit_fort[2] - 0.5:
            parts.append(f". Le critere {crit_faible[0]} ressort comme le moins bien evalue.")
        else:
            parts.append(".")
    else:
        parts.append(".")

    return " ".join(parts)


def _h1_doc(doc, text, color_hex="1F4E79"):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def _h2_doc(doc, text, color_hex="1F4E79"):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def _axe_banner_doc(doc, axe_num, label):
    color = AXE_COLORS[axe_num].lstrip("#")
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0); set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run(f"  AXE {axe_num} : {label.upper()}")
    r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
    doc.add_paragraph()


def export_synthese_docx(df):
    """
    Genere la synthese en .docx au format du Grille_RTE_OT_SudVienne.docx :
      1. Titre + metadonnees
      2. Section 5 : Grille - bloc par critere (score auto + elements + recommandations)
      3. Section 6 : Synthese des resultats (tableau Axe / Score / Niveau / Commentaire)
      4. Annexe : Checklists detaillees par critere
    """
    if len(df) == 0:
        # Doc minimal si pas de donnees
        doc = Document()
        add_para_styled(doc, "Aucune reponse a synthetiser pour l'instant.",
                        size=14, color_hex="9C2A2A", italic=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        buf = BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)
        s.top_margin = s.bottom_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Recuperation org/territoire (depuis la 1ere reponse)
    orga = df.iloc[0]["organisation"] if len(df) > 0 else "—"
    territoire = df.iloc[0]["territoire"] if len(df) > 0 else "—"

    # Calculs globaux
    agg = aggregate_responses(df)
    all_scores = []
    for _, row in df.iterrows():
        d = parse_donnees(row["donnees_json"])
        all_scores.append({a: score_axe(d, a) for a in AXES})
    df_scores = pd.DataFrame(all_scores)
    scores_moyens = {a: (df_scores[a].dropna().mean() if df_scores[a].dropna().shape[0] > 0 else None)
                     for a in AXES}
    valid = [s for s in scores_moyens.values() if s is not None]
    score_g = sum(valid) / len(valid) if valid else None

    # ============================================================
    # 1. TITRE + METADONNEES
    # ============================================================
    p = doc.add_paragraph()
    r = p.add_run("Grille de maturite RTE - Synthese collective")
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4E79")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run(orga); r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("2C3E50")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Modele V4 - Chaire TerrESS / Bordeaux Sciences Agro")
    r.font.size = Pt(11); r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("7D8B96")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadonnees
    repondants_list = ", ".join(
        f"{row['nom_repondant']}" + (f" ({row['fonction']})" if row['fonction'] else "")
        for _, row in df.iterrows()
    )
    meta = [
        ("Organisation evaluee", orga),
        ("Territoire concerne", territoire or "—"),
        ("Date de generation", date.today().strftime("%d / %m / %Y")),
        ("Nombre de repondants", str(len(df))),
        ("Repondants", repondants_list),
        ("Methode", "Notation automatique : Observe=1 / Partiel=0,5 / Non observe=0. "
                    "Score critere = (moyenne / nb questions) x 4."),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(12)
    for ri, (k, v) in enumerate(meta):
        c0 = tbl.cell(ri, 0); c1 = tbl.cell(ri, 1)
        set_cell_bg(c0, "DCE6F1")
        set_cell_border(c0); set_cell_border(c1)
        p = c0.paragraphs[0]
        r = p.add_run(k); r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("1F4E79")
        p = c1.paragraphs[0]
        r = p.add_run(v); r.font.size = Pt(10)

    # Note
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Synthese generee automatiquement a partir des reponses collectees "
        "via le formulaire en ligne. Cette synthese constitue un point de depart "
        "pour un atelier de co-evaluation avec les parties prenantes."
    )
    r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("7D8B96")

    # ============================================================
    # 2. SECTION 5 : GRILLE PAR CRITERE
    # ============================================================
    doc.add_page_break()
    _h1_doc(doc, "5. Grille de maturite : evaluation par critere")
    p = doc.add_paragraph()
    r = p.add_run(
        "Pour chaque critere : score automatique calcule a partir des reponses "
        "collectees, elements d'analyse (preuves citees par les repondants) "
        "et points a ameliorer."
    )
    r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("7D8B96")
    doc.add_paragraph()

    GRILLE_PRINCIPALE = [
        (1, "Diagnostic territorial", ["1.1", "1.2", "1.3"]),
        (2, "Vision territoriale et recits", ["2.1", "2.2", "2.3", "2.4", "2.5"]),
        (3, "Cooperation territoriale", ["3.1", "3.2"]),
        (4, "Gouvernance inclusive", ["4.1", "4.2", "4.3"]),
        (5, "Redistribuer et transformer", ["5.1", "5.2", "5.3", "5.4"]),
    ]
    titre_map = {c[0]: c[2] for c in CRITERES}

    for axe_num, axe_label, criteres_ids in GRILLE_PRINCIPALE:
        _axe_banner_doc(doc, axe_num, axe_label)
        color = AXE_COLORS[axe_num].lstrip("#")
        for c_id in criteres_ids:
            a = agg[c_id]
            score = a["score_moyen"]
            titre = titre_map[c_id]

            # Tableau 3 rangees (header, analyse, reco)
            tbl = doc.add_table(rows=3, cols=1)

            # Header
            c0 = tbl.cell(0, 0); set_cell_bg(c0, color); set_cell_border(c0, color)
            p = c0.paragraphs[0]
            r = p.add_run(f"  {c_id}  ")
            r.font.size = Pt(12); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            r = p.add_run(titre)
            r.font.size = Pt(12); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            if score is not None:
                r = p.add_run(
                    f"      Score : {score:.2f} / 4   |   Niveau : {qualif(score).split(' - ')[0]}   |   {a['nb_reponses']} reponses"
                )
            else:
                r = p.add_run("      Score : non evalue")
            r.font.size = Pt(11); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("FFE4A0")

            # Analyse : preuves agregees
            c1 = tbl.cell(1, 0); set_cell_bg(c1, "F4F6F8"); set_cell_border(c1)
            p = c1.paragraphs[0]
            r = p.add_run("Elements d'analyse (preuves citees par les repondants) : ")
            r.font.size = Pt(10); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(color)

            # On liste les preuves question par question
            preuves_total = 0
            for q in CHECKLISTS[c_id]:
                qa = a["questions"][q["id"]]
                if qa["preuves"]:
                    preuves_total += len(qa["preuves"])
                    p = c1.add_paragraph()
                    r = p.add_run(f"• {q['id']} : ")
                    r.font.size = Pt(9.5); r.font.bold = True
                    r = p.add_run(" ; ".join(qa["preuves"]))
                    r.font.size = Pt(9.5)
            if preuves_total == 0:
                p = c1.add_paragraph()
                r = p.add_run("Aucune preuve documentaire renseignee par les repondants.")
                r.font.size = Pt(10); r.font.italic = True
                r.font.color.rgb = RGBColor.from_string("7D8B96")

            # Reco : auto - basee sur questions faibles
            c2 = tbl.cell(2, 0); set_cell_bg(c2, "FFFAEB"); set_cell_border(c2)
            p = c2.paragraphs[0]
            r = p.add_run("Points a ameliorer (auto-deduits) : ")
            r.font.size = Pt(10); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("B7791F")

            questions_faibles = []
            for q in CHECKLISTS[c_id]:
                qa = a["questions"][q["id"]]
                total = qa["obs"] + qa["part"] + qa["non"]
                if total == 0:
                    continue
                q_score = (qa["obs"] * 1.0 + qa["part"] * 0.5) / total
                if q_score < 0.5:  # plus de "non" et "partiel" que "observe"
                    questions_faibles.append((q["id"], q["label"], qa))

            if questions_faibles:
                for q_id, q_label, qa in questions_faibles[:4]:
                    p = c2.add_paragraph()
                    r = p.add_run(f"• {q_id} : ")
                    r.font.size = Pt(9.5); r.font.bold = True
                    r = p.add_run(
                        f"{q_label} "
                        f"(✓{qa['obs']} | ~{qa['part']} | ✗{qa['non']})"
                    )
                    r.font.size = Pt(9.5); r.font.italic = True
            else:
                p = c2.add_paragraph()
                r = p.add_run("Aucun point d'amelioration prioritaire identifie sur ce critere.")
                r.font.size = Pt(10); r.font.italic = True
                r.font.color.rgb = RGBColor.from_string("2E8B73")

            doc.add_paragraph()
        doc.add_paragraph()

    # ============================================================
    # 3. SECTION 6 : SYNTHESE DES RESULTATS (FORMAT IDENTIQUE A L'EXISTANT)
    # ============================================================
    doc.add_page_break()
    _h1_doc(doc, "6. Synthese des resultats")

    tbl = doc.add_table(rows=len(AXES) + 2, cols=4)
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(2.5)
    tbl.columns[2].width = Cm(3.5)
    tbl.columns[3].width = Cm(5)

    # Header
    for ci, h in enumerate(["Axe", "Score moyen", "Niveau", "Commentaire synthetique"]):
        cell = tbl.cell(0, ci); set_cell_bg(cell, "1F4E79"); set_cell_border(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h); r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lignes axes
    for ri, axe_num in enumerate(AXES, start=1):
        color = AXE_COLORS[axe_num].lstrip("#")
        score = scores_moyens[axe_num]
        nb_crit = NB_CRITERES_PAR_AXE[axe_num]
        commentaire = _commentaire_axe(axe_num, score, scores_moyens, agg, df)

        # Col 1 : Axe (colore)
        cell = tbl.cell(ri, 0); set_cell_bg(cell, color); set_cell_border(cell)
        p = cell.paragraphs[0]
        r = p.add_run(f"Axe {axe_num} - {AXES[axe_num]}")
        r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        p2 = cell.add_paragraph()
        r = p2.add_run(f"({nb_crit} criteres)")
        r.font.size = Pt(9); r.font.italic = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")

        # Col 2 : Score
        cell = tbl.cell(ri, 1); set_cell_bg(cell, "F4F6F8"); set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{score:.2f} / 4" if score is not None else "—")
        r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(color)

        # Col 3 : Niveau
        cell = tbl.cell(ri, 2); set_cell_bg(cell, "F4F6F8"); set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(qualif(score))
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("2C3E50")

        # Col 4 : Commentaire
        cell = tbl.cell(ri, 3); set_cell_bg(cell, "F4F6F8"); set_cell_border(cell)
        p = cell.paragraphs[0]
        r = p.add_run(commentaire); r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("2C3E50")

    # Ligne globale doree
    ri = len(AXES) + 1
    for ci, val in enumerate([
        "SCORE GLOBAL RTE",
        f"{score_g:.2f} / 4" if score_g is not None else "—",
        qualif(score_g),
        f"Moyenne ponderee des 5 axes ({TOTAL_CRITERES} criteres au total)",
    ]):
        cell = tbl.cell(ri, ci); set_cell_bg(cell, "C9A961"); set_cell_border(cell)
        p = cell.paragraphs[0]
        r = p.add_run(val); r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string("2C3E50")
        if ci in (1, 2):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lecture d'ensemble
    doc.add_paragraph()
    _h2_doc(doc, "Lecture d'ensemble")
    if score_g is not None:
        p = doc.add_paragraph()
        r = p.add_run(
            f"Avec un score global de {score_g:.2f}/4, l'organisation evaluee se situe au niveau "
            f"\"{qualif(score_g)}\" sur la base des {len(df)} reponses collectees."
        )
        r.font.size = Pt(11)

        # Points forts
        axes_tries = sorted(
            ((a, s) for a, s in scores_moyens.items() if s is not None),
            key=lambda x: x[1], reverse=True,
        )
        if axes_tries:
            doc.add_paragraph()
            r = doc.add_paragraph().add_run("Points forts :")
            r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor.from_string("2E8B73")
            for axe_num, sc in axes_tries[:2]:
                if sc >= 2.5:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(
                        f"Axe {axe_num} - {AXES[axe_num]} ({sc:.2f}/4) : "
                        f"{qualif(sc).lower()}."
                    )
                    r.font.size = Pt(10.5)

            # Axes a renforcer
            r = doc.add_paragraph().add_run("Axes a renforcer :")
            r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor.from_string("9C2A2A")
            for axe_num, sc in axes_tries[::-1][:2]:
                if sc < 3.0:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(
                        f"Axe {axe_num} - {AXES[axe_num]} ({sc:.2f}/4) : "
                        f"{qualif(sc).lower()}."
                    )
                    r.font.size = Pt(10.5)

    # ============================================================
    # 4. ANNEXE : DETAIL DES REPONSES PAR CRITERE
    # ============================================================
    doc.add_page_break()
    _h1_doc(doc, "Annexe - Detail des reponses par critere")
    p = doc.add_paragraph()
    r = p.add_run(
        "Pour chaque critere : distribution des reponses question par question "
        "et preuves collectees. ✓ Observe = 1 pt | ~ Partiel = 0,5 pt | ✗ Non observe = 0 pt."
    )
    r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("7D8B96")

    for axe_num, axe_label, criteres_ids in GRILLE_PRINCIPALE:
        _axe_banner_doc(doc, axe_num, axe_label)
        color = AXE_COLORS[axe_num].lstrip("#")
        for c_id in criteres_ids:
            a = agg[c_id]
            titre = titre_map[c_id]

            # Bandeau critere
            tbl = doc.add_table(rows=1, cols=1)
            cell = tbl.cell(0, 0); set_cell_bg(cell, color)
            p = cell.paragraphs[0]
            r = p.add_run(f"  {c_id}  ")
            r.font.size = Pt(11); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            r = p.add_run(titre)
            r.font.size = Pt(11); r.font.bold = True
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            if a["score_moyen"] is not None:
                r = p.add_run(
                    f"   ·   Score : {a['score_moyen']:.2f}/4   ·   {a['nb_reponses']} reponses"
                )
                r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("FFE4A0")

            # Tableau questions detail
            questions = CHECKLISTS[c_id]
            tbl = doc.add_table(rows=len(questions) + 1, cols=4)
            tbl.columns[0].width = Cm(7.5)
            tbl.columns[1].width = Cm(3)
            tbl.columns[2].width = Cm(5)
            tbl.columns[3].width = Cm(1.5)

            for ci, h in enumerate(["Action a documenter", "Distribution",
                                    "Preuves collectees", "Score q."]):
                cell = tbl.cell(0, ci); set_cell_bg(cell, "DCE6F1"); set_cell_border(cell)
                p = cell.paragraphs[0]
                r = p.add_run(h); r.font.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = RGBColor.from_string("1F4E79")

            for ri, q in enumerate(questions, start=1):
                qa = a["questions"][q["id"]]
                total = qa["obs"] + qa["part"] + qa["non"]
                q_score = ((qa["obs"] * 1.0 + qa["part"] * 0.5) / total * 4) if total > 0 else None

                # Question
                c0 = tbl.cell(ri, 0); set_cell_border(c0); c0.text = ""
                p = c0.paragraphs[0]
                r = p.add_run(f"{q['id']} "); r.font.bold = True; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor.from_string("1F4E79")
                r = p.add_run(q["label"]); r.font.size = Pt(9.5)

                # Distribution
                c1 = tbl.cell(ri, 1); set_cell_border(c1); c1.text = ""
                p = c1.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"✓ {qa['obs']}")
                r.font.size = Pt(9.5); r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("2E8B73")
                r = p.add_run("  |  "); r.font.size = Pt(9.5)
                r = p.add_run(f"~ {qa['part']}")
                r.font.size = Pt(9.5); r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("C9A961")
                r = p.add_run("  |  "); r.font.size = Pt(9.5)
                r = p.add_run(f"✗ {qa['non']}")
                r.font.size = Pt(9.5); r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("C0392B")

                # Preuves
                c2 = tbl.cell(ri, 2); set_cell_border(c2); c2.text = ""
                if qa["preuves"]:
                    p = c2.paragraphs[0]
                    r = p.add_run(qa["preuves"][0])
                    r.font.size = Pt(9); r.font.italic = True
                    for preuve in qa["preuves"][1:]:
                        p = c2.add_paragraph()
                        r = p.add_run(preuve)
                        r.font.size = Pt(9); r.font.italic = True
                else:
                    p = c2.paragraphs[0]
                    r = p.add_run("—")
                    r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("BFBFBF")

                # Score q.
                c3 = tbl.cell(ri, 3); set_cell_border(c3); c3.text = ""
                p = c3.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if q_score is not None:
                    r = p.add_run(f"{q_score:.2f}")
                    r.font.size = Pt(10); r.font.bold = True
                else:
                    r = p.add_run("—"); r.font.size = Pt(9.5)
            doc.add_paragraph()

    # Note finale
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        f"--- Synthese generee automatiquement le {date.today().strftime('%d/%m/%Y')} ---"
    )
    r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("7D8B96")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# UI : CSS / STYLE
# ============================================================

def apply_custom_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Crimson+Pro:wght@600;700&display=swap');

    /* === BASE === */
    html, body, .stApp, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    }
    .stApp {
        background: linear-gradient(180deg, #F7F8FC 0%, #EEF2F8 100%);
        background-attachment: fixed;
    }
    .block-container {
        padding-top: 2rem !important;
        max-width: 1200px;
    }

    /* === TYPO === */
    h1, h2, h3 {
        font-family: 'Inter', sans-serif !important;
        color: #0F2A47 !important;
        letter-spacing: -0.02em !important;
    }
    h1 { font-weight: 800 !important; }
    h2 { font-weight: 700 !important; }
    h3 { font-weight: 600 !important; }

    /* === HERO === */
    .hero {
        background: linear-gradient(135deg, #0F2A47 0%, #1F4E79 50%, #2A6AAA 100%);
        padding: 36px 40px;
        border-radius: 20px;
        margin-bottom: 28px;
        color: white;
        position: relative;
        overflow: hidden;
        box-shadow: 0 20px 50px -20px rgba(15, 42, 71, 0.4);
    }
    .hero::before {
        content: "";
        position: absolute;
        top: -50%;
        right: -10%;
        width: 400px;
        height: 400px;
        background: radial-gradient(circle, rgba(225, 182, 96, 0.25) 0%, transparent 70%);
        border-radius: 50%;
    }
    .hero-tag {
        display: inline-block;
        background: rgba(225, 182, 96, 0.2);
        color: #E1B660;
        padding: 5px 14px;
        border-radius: 20px;
        font-size: 11px;
        font-weight: 600;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        margin-bottom: 14px;
        border: 1px solid rgba(225, 182, 96, 0.3);
    }
    .hero-title {
        font-size: 36px;
        font-weight: 800;
        margin: 0 0 8px 0;
        letter-spacing: -0.03em;
        line-height: 1.1;
    }
    .hero-subtitle {
        font-size: 16px;
        opacity: 0.85;
        font-weight: 400;
        margin: 0;
        max-width: 700px;
        line-height: 1.6;
    }
    .hero-meta {
        margin-top: 20px;
        display: flex;
        gap: 32px;
        flex-wrap: wrap;
    }
    .hero-meta-item {
        display: flex;
        flex-direction: column;
    }
    .hero-meta-val {
        font-size: 24px;
        font-weight: 700;
        color: #E1B660;
    }
    .hero-meta-lbl {
        font-size: 11px;
        opacity: 0.8;
        text-transform: uppercase;
        letter-spacing: 0.08em;
    }

    /* === AXE BANNER === */
    .axe-banner {
        padding: 14px 22px;
        border-radius: 12px;
        margin: 32px 0 18px 0;
        color: white;
        font-weight: 700;
        font-size: 16px;
        letter-spacing: 0.02em;
        display: flex;
        align-items: center;
        gap: 14px;
        box-shadow: 0 6px 16px -4px rgba(0, 0, 0, 0.15);
    }
    .axe-banner .axe-num {
        background: rgba(255, 255, 255, 0.25);
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        backdrop-filter: blur(4px);
    }

    /* === CRITERE CARD === */
    .crit-card {
        background: #FFFFFF;
        border-radius: 14px;
        padding: 20px 24px;
        margin-bottom: 14px;
        border: 1px solid #E5E9F0;
        box-shadow: 0 1px 3px rgba(15, 42, 71, 0.04);
        transition: box-shadow 0.2s;
    }
    .crit-card:hover {
        box-shadow: 0 8px 24px -8px rgba(15, 42, 71, 0.12);
    }
    .crit-head {
        display: flex;
        align-items: center;
        gap: 12px;
        margin-bottom: 6px;
    }
    .crit-id {
        background: #0F2A47;
        color: #FFF;
        padding: 4px 10px;
        border-radius: 6px;
        font-weight: 700;
        font-size: 12px;
        letter-spacing: 0.04em;
    }
    .crit-title {
        font-size: 16px;
        font-weight: 600;
        color: #0F2A47;
    }
    .crit-hint {
        font-size: 12px;
        color: #7D8B96;
        margin: 8px 0 14px 0;
        padding: 8px 12px;
        background: #F7F8FC;
        border-left: 3px solid #E1B660;
        border-radius: 4px;
    }

    /* === KPI === */
    .kpi {
        background: white;
        padding: 22px 18px;
        border-radius: 14px;
        text-align: center;
        box-shadow: 0 2px 6px rgba(15, 42, 71, 0.06);
        border: 1px solid #E5E9F0;
        transition: transform 0.15s, box-shadow 0.15s;
    }
    .kpi:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 24px -6px rgba(15, 42, 71, 0.15);
    }
    .kpi-val {
        font-size: 36px;
        font-weight: 800;
        color: #0F2A47;
        line-height: 1;
        letter-spacing: -0.03em;
    }
    .kpi-lbl {
        font-size: 11px;
        color: #7D8B96;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin-top: 8px;
        font-weight: 600;
    }
    .kpi-accent { color: #E1B660; }
    .kpi-success { color: #2D7A65; }
    .kpi-danger { color: #B53B3B; }

    /* === RADIO BUTTONS === */
    div[data-testid="stRadio"] > label > div > p {
        font-size: 13px;
        color: #2C3E50;
    }
    div[data-testid="stRadio"] label {
        background: transparent;
        margin-bottom: 2px !important;
    }

    /* === TEXT INPUTS === */
    .stTextInput input, .stTextArea textarea {
        border-radius: 10px !important;
        border: 1px solid #D7DEE8 !important;
        background: #FFFFFF !important;
        font-size: 13px !important;
        transition: border-color 0.15s, box-shadow 0.15s !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: #1F4E79 !important;
        box-shadow: 0 0 0 3px rgba(31, 78, 121, 0.12) !important;
    }

    /* === BUTTONS === */
    .stButton button, .stDownloadButton button, .stFormSubmitButton button {
        border-radius: 10px !important;
        font-weight: 600 !important;
        letter-spacing: 0.02em !important;
        transition: all 0.15s !important;
        border: none !important;
    }
    .stButton button[kind="primary"], .stDownloadButton button[kind="primary"],
    .stFormSubmitButton button {
        background: linear-gradient(135deg, #1F4E79 0%, #2A6AAA 100%) !important;
        color: white !important;
        box-shadow: 0 4px 12px -2px rgba(31, 78, 121, 0.35) !important;
    }
    .stButton button[kind="primary"]:hover, .stDownloadButton button[kind="primary"]:hover,
    .stFormSubmitButton button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 8px 20px -4px rgba(31, 78, 121, 0.45) !important;
    }

    /* === SIDEBAR === */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0F2A47 0%, #1A3A5C 100%) !important;
    }
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] [data-testid="stMetricLabel"],
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] p {
        color: #E5E9F0 !important;
    }
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        color: #FFFFFF !important;
    }
    section[data-testid="stSidebar"] [data-testid="stMetricValue"] {
        color: #E1B660 !important;
        font-weight: 800;
    }
    section[data-testid="stSidebar"] hr {
        border-color: rgba(255, 255, 255, 0.12) !important;
    }
    section[data-testid="stSidebar"] div[data-testid="stRadio"] label p {
        color: #E5E9F0 !important;
        font-size: 14px;
    }

    /* === PROGRESS BAR === */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #1F4E79 0%, #E1B660 100%) !important;
    }

    /* === TABLES === */
    .stDataFrame {
        border-radius: 12px !important;
        overflow: hidden !important;
        box-shadow: 0 1px 3px rgba(15, 42, 71, 0.06);
    }

    /* === EXPANDERS === */
    div[data-testid="stExpander"] {
        background: #FFFFFF;
        border: 1px solid #E5E9F0 !important;
        border-radius: 12px !important;
        overflow: hidden;
        box-shadow: 0 1px 3px rgba(15, 42, 71, 0.04);
        margin-bottom: 12px;
    }
    div[data-testid="stExpander"] summary {
        font-weight: 600 !important;
        color: #0F2A47 !important;
        padding: 14px 18px !important;
    }
    div[data-testid="stExpander"] summary:hover {
        background: #F7F8FC;
    }

    /* === ALERTS === */
    div[data-testid="stAlert"] {
        border-radius: 12px !important;
        border: none !important;
    }

    /* === SCORE BADGE === */
    .score-badge {
        display: inline-block;
        background: linear-gradient(135deg, #1F4E79 0%, #2A6AAA 100%);
        color: white;
        padding: 8px 16px;
        border-radius: 10px;
        font-weight: 700;
        font-size: 14px;
        letter-spacing: 0.02em;
        box-shadow: 0 3px 8px -2px rgba(31, 78, 121, 0.3);
    }

    /* === QUESTION ROW === */
    .question-block {
        background: #FAFBFD;
        border: 1px solid #EAEFF5;
        border-radius: 10px;
        padding: 12px 14px;
        margin-bottom: 10px;
    }
    .q-id-tag {
        display: inline-block;
        background: #0F2A47;
        color: #FFF;
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 10px;
        font-weight: 700;
        letter-spacing: 0.04em;
        margin-right: 8px;
        font-family: 'Inter', monospace;
    }
    .q-text {
        color: #2C3E50;
        font-size: 13.5px;
        line-height: 1.5;
    }

    /* === DIVIDER === */
    .section-divider {
        text-align: center;
        margin: 32px 0 18px 0;
        position: relative;
    }
    .section-divider::before {
        content: "";
        position: absolute;
        top: 50%;
        left: 0; right: 0;
        height: 1px;
        background: linear-gradient(90deg, transparent, #D7DEE8, transparent);
    }
    .section-divider span {
        background: #F7F8FC;
        padding: 0 18px;
        color: #7D8B96;
        font-size: 11px;
        text-transform: uppercase;
        letter-spacing: 0.15em;
        font-weight: 600;
        position: relative;
    }
    </style>
    """, unsafe_allow_html=True)


# ============================================================
# UI : PAGE FORMULAIRE
# ============================================================

def page_formulaire():
    # Hero header
    nb_questions_total = sum(len(items) for items in CHECKLISTS.values())
    st.markdown(f"""
    <div class='hero'>
        <div class='hero-tag'>Grille de maturite RTE · Modele V4</div>
        <h1 class='hero-title'>Evaluation collaborative de la maturite RTE</h1>
        <p class='hero-subtitle'>
            Pour chaque question, indiquez si la pratique est <b>Observee</b>,
            <b>Partielle</b> ou <b>Non observee</b>, et ajoutez une preuve quand vous pouvez.
            La notation se calcule automatiquement.
        </p>
        <div class='hero-meta'>
            <div class='hero-meta-item'>
                <span class='hero-meta-val'>5</span>
                <span class='hero-meta-lbl'>Axes</span>
            </div>
            <div class='hero-meta-item'>
                <span class='hero-meta-val'>{TOTAL_CRITERES}</span>
                <span class='hero-meta-lbl'>Criteres</span>
            </div>
            <div class='hero-meta-item'>
                <span class='hero-meta-val'>{nb_questions_total}</span>
                <span class='hero-meta-lbl'>Questions</span>
            </div>
            <div class='hero-meta-item'>
                <span class='hero-meta-val'>15-30</span>
                <span class='hero-meta-lbl'>Minutes</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Mode d'emploi notation
    cinf1, cinf2, cinf3 = st.columns(3)
    with cinf1:
        st.markdown(
            "<div class='kpi'><div class='kpi-val kpi-success'>1 pt</div>"
            "<div class='kpi-lbl'>✅ Observe</div></div>",
            unsafe_allow_html=True)
    with cinf2:
        st.markdown(
            "<div class='kpi'><div class='kpi-val kpi-accent'>0,5 pt</div>"
            "<div class='kpi-lbl'>🟡 Partiel</div></div>",
            unsafe_allow_html=True)
    with cinf3:
        st.markdown(
            "<div class='kpi'><div class='kpi-val kpi-danger'>0 pt</div>"
            "<div class='kpi-lbl'>❌ Non observe</div></div>",
            unsafe_allow_html=True)

    st.markdown("<div style='height:24px'></div>", unsafe_allow_html=True)

    # Statut options
    statut_options = [None, "obs", "part", "non"]
    statut_format = {
        None: "— Pas evalue",
        "obs": "✅ Observe",
        "part": "🟡 Partiel",
        "non": "❌ Non observe",
    }

    # ========================================================================
    # TOUT LE FORMULAIRE DANS st.form -> aucun rechargement = aucun saut
    # ========================================================================
    with st.form(key="rte_form", clear_on_submit=False, border=False):

        # --- Coordonnees ---
        st.markdown("""
        <div class='section-divider'><span>VOS COORDONNEES</span></div>
        """, unsafe_allow_html=True)
        st.markdown(
            "<p style='color:#7D8B96;font-size:13px;margin-bottom:18px'>"
            "Les champs marques d'un <b style='color:#9C2A2A'>*</b> sont obligatoires."
            "</p>",
            unsafe_allow_html=True,
        )

        col1, col2 = st.columns(2)
        with col1:
            st.text_input(
                "Nom et prenom *", placeholder="ex. Jean Dupont",
                key="meta_nom",
            )
            st.text_input(
                "Email *", placeholder="ex. jean.dupont@example.com",
                key="meta_email",
                help="Pour vous tenir informe et pouvoir vous recontacter si besoin.",
            )
            st.text_input(
                "Fonction", placeholder="ex. Directeur, charge de mission",
                key="meta_fonction",
            )
        with col2:
            st.text_input(
                "Organisation evaluee *",
                value="",
                placeholder="ex. Office de Tourisme Sud Vienne Poitou",
                key="meta_orga",
                help="Nom exact de l'organisation que vous evaluez. "
                     "Les reponses portant le meme nom seront regroupees dans le dashboard.",
            )
            st.text_input(
                "Territoire concerne",
                value="",
                placeholder="ex. Communaute de Communes Vienne et Gartempe",
                key="meta_territoire",
            )

        # --- Boucle sur les axes ---
        for axe_num, axe_label in AXES.items():
            color = AXE_COLORS[axe_num]
            criteres_axe = [(c, a, t) for c, a, t in CRITERES if a == axe_num]
            nb_q_axe = sum(len(CHECKLISTS[c_id]) for c_id, _, _ in criteres_axe)

            st.markdown(
                f"<div class='axe-banner' style='background:linear-gradient(135deg,{color} 0%, {color}DD 100%)'>"
                f"<span class='axe-num'>AXE {axe_num}</span>"
                f"<span>{axe_label.upper()}</span>"
                f"<span style='margin-left:auto;font-size:12px;opacity:0.85;font-weight:500'>"
                f"{len(criteres_axe)} criteres · {nb_q_axe} questions</span>"
                f"</div>",
                unsafe_allow_html=True
            )

            for c_id, _, titre in criteres_axe:
                # Carte critere
                st.markdown(
                    f"<div class='crit-card'>"
                    f"<div class='crit-head'>"
                    f"<span class='crit-id'>{c_id}</span>"
                    f"<span class='crit-title'>{titre}</span>"
                    f"</div>"
                    f"<div class='crit-hint'>💡 {len(CHECKLISTS[c_id])} questions a evaluer. "
                    f"Score auto = (moyenne / nb questions) × 4.</div>"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Questions
                for q in CHECKLISTS[c_id]:
                    st.markdown(
                        f"<div class='question-block'>"
                        f"<span class='q-id-tag'>{q['id']}</span>"
                        f"<span class='q-text'>{q['label']}</span>"
                        f"</div>",
                        unsafe_allow_html=True
                    )
                    cstat, cpr = st.columns([1, 1.5])
                    with cstat:
                        st.radio(
                            label=f"Statut {q['id']}",
                            options=statut_options,
                            format_func=lambda x: statut_format[x],
                            key=f"st_{q['id']}",
                            label_visibility="collapsed",
                            horizontal=False,
                            index=0,
                        )
                    with cpr:
                        st.text_area(
                            label=f"Preuve {q['id']}",
                            key=f"pr_{q['id']}",
                            placeholder="Document, exemple, source, indicateur...",
                            label_visibility="collapsed",
                            height=85,
                        )

                st.text_area(
                    "Commentaire global sur ce critere (facultatif)",
                    key=f"com_{c_id}",
                    placeholder="Remarques transversales, nuances, contexte...",
                    height=70,
                )
                st.markdown("<div style='height:18px'></div>", unsafe_allow_html=True)

        # --- Commentaire global ---
        st.markdown("""
        <div class='section-divider'><span>POUR FINIR</span></div>
        """, unsafe_allow_html=True)
        st.text_area(
            "Commentaire global / remarques transversales",
            placeholder="(optionnel) Remarques qui ne rentrent pas dans la grille...",
            height=100,
            key="meta_com_global",
        )

        # --- Submit button ---
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        submitted = st.form_submit_button(
            "✅ Soumettre ma reponse",
            type="primary", use_container_width=True,
        )

    # ========================================================================
    # APRES SUBMIT : lecture des valeurs et sauvegarde
    # ========================================================================
    if submitted:
        # Reconstitue les donnees depuis session_state
        nom = st.session_state.get("meta_nom", "").strip()
        fonction = st.session_state.get("meta_fonction", "").strip()
        orga = st.session_state.get("meta_orga", "").strip()
        territoire = st.session_state.get("meta_territoire", "").strip()
        email = st.session_state.get("meta_email", "").strip()
        commentaire_global = st.session_state.get("meta_com_global", "").strip()

        criteres_data = {}
        for c_id, _, _ in CRITERES:
            questions = {}
            for q in CHECKLISTS[c_id]:
                questions[q["id"]] = {
                    "statut": st.session_state.get(f"st_{q['id']}"),
                    "preuve": st.session_state.get(f"pr_{q['id']}", "").strip(),
                }
            criteres_data[c_id] = {
                "questions": questions,
                "commentaire": st.session_state.get(f"com_{c_id}", "").strip(),
            }

        # Validation
        nb_criteres_repondus = sum(
            1 for c_id, _, _ in CRITERES
            if score_critere(criteres_data[c_id]["questions"]) is not None
        )

        if not nom:
            st.error("⚠️ Merci d'indiquer votre nom et prenom.")
        elif not email or "@" not in email or "." not in email.split("@")[-1]:
            st.error("⚠️ Merci d'indiquer un email valide.")
        elif not orga:
            st.error("⚠️ Merci d'indiquer l'organisation evaluee.")
        elif nb_criteres_repondus == 0:
            st.error("⚠️ Merci d'evaluer au moins une question avant de soumettre.")
        else:
            save_response(
                meta={
                    "nom": nom, "fonction": fonction, "organisation": orga,
                    "territoire": territoire, "email": email,
                    "commentaire_global": commentaire_global,
                },
                criteres_data=criteres_data,
            )
            score_g = score_global(criteres_data)
            st.success(
                f"🎉 Merci **{nom}** ! Votre reponse a bien ete enregistree.\n\n"
                f"**Votre score global** : {score_g:.2f} / 4 — {qualif(score_g)}\n\n"
                f"Criteres evalues : {nb_criteres_repondus} / {TOTAL_CRITERES}"
            )
            st.balloons()


# ============================================================
# UI : PAGE DASHBOARD
# ============================================================

def page_dashboard():
    st.title("📊 Dashboard - Synthese collective en temps reel")

    df_all = load_responses()
    if len(df_all) == 0:
        st.info("👋 Aucune reponse pour l'instant. Partagez le lien du formulaire avec vos collegues !")
        return

    # ============================================================
    # FILTRE PAR ORGANISATION
    # ============================================================
    organisations_uniques = sorted(df_all["organisation"].dropna().unique().tolist())
    OPTION_TOUTES = "🌐 Toutes les organisations"
    options_filtre = [OPTION_TOUTES] + organisations_uniques

    st.markdown(
        "<div style='background:linear-gradient(135deg,#0F2A47 0%,#1F4E79 100%);"
        "padding:18px 24px;border-radius:14px;margin-bottom:20px;color:white;'>"
        "<div style='font-size:11px;font-weight:600;letter-spacing:0.08em;"
        "color:#E1B660;text-transform:uppercase;margin-bottom:4px'>"
        "FILTRER PAR ORGANISATION EVALUEE</div>"
        "<div style='opacity:0.85;font-size:13px'>Choisissez une organisation pour voir "
        "uniquement ses repondants et ses resultats.</div>"
        "</div>",
        unsafe_allow_html=True,
    )

    filtre_orga = st.selectbox(
        "Organisation",
        options=options_filtre,
        index=0,
        label_visibility="collapsed",
    )

    if filtre_orga == OPTION_TOUTES:
        df = df_all
        org_filtree = None
    else:
        df = df_all[df_all["organisation"] == filtre_orga].reset_index(drop=True)
        org_filtree = filtre_orga

    if len(df) == 0:
        st.warning(f"Aucune reponse pour l'organisation : {org_filtree}")
        return

    # ============================================================
    # BLOC ORGANISATION + REPONDANTS (visible si org filtree)
    # ============================================================
    if org_filtree:
        repondants_blocks = []
        for _, row in df.iterrows():
            nom = row["nom_repondant"]
            fonction = row["fonction"] or ""
            email = row["email"] or ""
            date_sub = row["date_soumission"][:10]
            fonction_html = (f" · <span style='color:#7D8B96;font-size:13px'>{fonction}</span>"
                             if fonction else "")
            email_html = f"📧 {email} · " if email else ""
            repondants_blocks.append(
                f"<div style='padding:10px 14px;background:#F7F8FC;border-radius:8px;"
                f"margin-bottom:6px;border-left:3px solid #1F4E79'>"
                f"<b style='color:#0F2A47'>{nom}</b>{fonction_html}"
                f"<br><span style='font-size:12px;color:#7D8B96'>"
                f"{email_html}📅 {date_sub}"
                f"</span></div>"
            )
        repondants_html = "".join(repondants_blocks)
        st.markdown(
            f"<div style='background:white;border-radius:14px;padding:24px;"
            f"box-shadow:0 2px 8px rgba(15,42,71,0.06);border:1px solid #E5E9F0;"
            f"margin-bottom:24px'>"
            f"<div style='display:flex;align-items:center;gap:12px;margin-bottom:16px'>"
            f"<div style='background:#0F2A47;color:white;padding:6px 14px;"
            f"border-radius:8px;font-size:11px;font-weight:700;letter-spacing:0.08em'>"
            f"ORGANISATION EVALUEE</div>"
            f"<div style='font-size:20px;font-weight:700;color:#0F2A47'>{org_filtree}</div>"
            f"</div>"
            f"<div style='color:#7D8B96;font-size:13px;margin-bottom:14px'>"
            f"<b>{len(df)} repondant{'s' if len(df) > 1 else ''}</b> "
            f"sur cette organisation</div>"
            f"{repondants_html}"
            f"</div>",
            unsafe_allow_html=True,
        )
    else:
        # Vue globale : nombre d'organisations + nombre de repondants par organisation
        nb_orgs = len(organisations_uniques)
        st.markdown(
            f"<div style='background:white;border-radius:14px;padding:24px;"
            f"box-shadow:0 2px 8px rgba(15,42,71,0.06);border:1px solid #E5E9F0;"
            f"margin-bottom:24px'>"
            f"<div style='display:flex;align-items:center;gap:12px;margin-bottom:16px'>"
            f"<div style='background:#E1B660;color:#0F2A47;padding:6px 14px;"
            f"border-radius:8px;font-size:11px;font-weight:700;letter-spacing:0.08em'>"
            f"VUE GLOBALE</div>"
            f"<div style='font-size:18px;font-weight:700;color:#0F2A47'>"
            f"{nb_orgs} organisation{'s' if nb_orgs > 1 else ''} evaluee{'s' if nb_orgs > 1 else ''} · "
            f"{len(df_all)} reponse{'s' if len(df_all) > 1 else ''} au total</div>"
            f"</div>"
            f"<div style='color:#7D8B96;font-size:13px'>"
            f"Selectionnez une organisation ci-dessus pour voir le detail des repondants et resultats."
            f"</div>"
            f"</div>",
            unsafe_allow_html=True,
        )

        # Mini-tableau de repartition par organisation
        repartition = (df_all.groupby("organisation")
                       .agg(nb_reponses=("id", "count"),
                            repondants=("nom_repondant", lambda x: ", ".join(sorted(x.unique()))))
                       .reset_index()
                       .sort_values("nb_reponses", ascending=False))
        repartition.columns = ["Organisation", "Nb reponses", "Repondants"]
        st.markdown("**Repartition par organisation :**")
        st.dataframe(repartition, hide_index=True, use_container_width=True)

    # ============================================================
    # SUITE DU DASHBOARD : calculs sur df (filtre)
    # ============================================================
    # Scores par axe
    all_scores = []
    for _, row in df.iterrows():
        d = parse_donnees(row["donnees_json"])
        all_scores.append({a: score_axe(d, a) for a in AXES})
    df_scores = pd.DataFrame(all_scores)

    def mean_safe(serie):
        valid = serie.dropna()
        return valid.mean() if len(valid) > 0 else None

    scores_moyens = {a: mean_safe(df_scores[a]) for a in AXES}
    valid_scores = [s for s in scores_moyens.values() if s is not None]
    score_g = sum(valid_scores) / len(valid_scores) if valid_scores else None

    # KPI
    st.subheader("Vue d'ensemble")
    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(
        f"<div class='kpi'><div class='kpi-val'>{len(df)}</div>"
        f"<div class='kpi-lbl'>Reponses</div></div>",
        unsafe_allow_html=True)
    c2.markdown(
        f"<div class='kpi'><div class='kpi-val'>"
        f"{score_g:.2f}" if score_g is not None else "—"
        f"</div><div class='kpi-lbl'>Score moyen / 4</div></div>",
        unsafe_allow_html=True)
    c3.markdown(
        f"<div class='kpi'><div class='kpi-val'>"
        f"{qualif(score_g).split(' - ')[0]}"
        f"</div><div class='kpi-lbl'>Niveau moyen</div></div>",
        unsafe_allow_html=True)
    nb_axes_repondus = sum(1 for s in scores_moyens.values() if s is not None)
    c4.markdown(
        f"<div class='kpi'><div class='kpi-val'>{nb_axes_repondus}/5</div>"
        f"<div class='kpi-lbl'>Axes evalues</div></div>",
        unsafe_allow_html=True)

    st.markdown("---")
    col1, col2 = st.columns(2)

    # Radar moyen
    with col1:
        st.subheader("Profil RTE moyen (radar)")
        valid_axes = [a for a in AXES if scores_moyens[a] is not None]
        if valid_axes:
            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(
                r=[scores_moyens[a] for a in valid_axes],
                theta=[f"Axe {a}<br>{AXES[a][:25]}" for a in valid_axes],
                fill="toself", name="Moyenne",
                line=dict(color="#1F4E79", width=3),
                fillcolor="rgba(31, 78, 121, 0.3)",
            ))
            fig.update_layout(
                polar=dict(
                    radialaxis=dict(visible=True, range=[0, 4], tickfont=dict(size=10)),
                    angularaxis=dict(tickfont=dict(size=11)),
                ),
                showlegend=False, height=420, margin=dict(l=60, r=60, t=20, b=20),
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Pas assez de donnees pour le radar.")

    # Barres horizontales
    with col2:
        st.subheader("Scores moyens par axe")
        valid_axes = [a for a in AXES if scores_moyens[a] is not None]
        if valid_axes:
            bar_data = pd.DataFrame({
                "Axe": [f"Axe {a} - {AXES[a]}" for a in valid_axes],
                "Score": [scores_moyens[a] for a in valid_axes],
            })
            fig = px.bar(
                bar_data, x="Score", y="Axe", orientation="h",
                color="Axe", color_discrete_sequence=[AXE_COLORS[a] for a in valid_axes],
                text="Score",
            )
            fig.update_traces(texttemplate="%{text:.2f}", textposition="outside")
            fig.update_layout(
                xaxis_range=[0, 4.5], showlegend=False, height=420,
                margin=dict(l=20, r=20, t=20, b=20),
                xaxis_title="Score moyen / 4", yaxis_title=None,
            )
            st.plotly_chart(fig, use_container_width=True)

    # Tableau scores par critere
    st.markdown("---")
    st.subheader("Scores par critere")
    agg = aggregate_responses(df)
    rows = []
    for c_id, axe, titre in CRITERES:
        a = agg[c_id]
        rows.append({
            "Critere": c_id,
            "Axe": f"Axe {axe}",
            "Titre": titre,
            "Score moyen / 4": f"{a['score_moyen']:.2f}" if a["score_moyen"] is not None else "—",
            "Min": f"{a['score_min']:.2f}" if a["score_min"] is not None else "—",
            "Max": f"{a['score_max']:.2f}" if a["score_max"] is not None else "—",
            "Nb reponses": a["nb_reponses"],
        })
    st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)

    # Comparaison entre repondants
    if len(df) > 1:
        st.markdown("---")
        st.subheader("Comparaison entre repondants")
        comparison = []
        for _, row in df.iterrows():
            d = parse_donnees(row["donnees_json"])
            entry = {"Repondant": row["nom_repondant"]}
            for a in AXES:
                s = score_axe(d, a)
                entry[f"Axe {a}"] = f"{s:.2f}" if s is not None else "—"
            sg = score_global(d)
            entry["Global"] = f"{sg:.2f}" if sg is not None else "—"
            comparison.append(entry)
        st.dataframe(pd.DataFrame(comparison), hide_index=True, use_container_width=True)

        # Radar comparatif
        st.subheader("Radars superposes")
        fig = go.Figure()
        for _, row in df.iterrows():
            d = parse_donnees(row["donnees_json"])
            scores = [score_axe(d, a) for a in AXES]
            # remplace les None par 0 pour l'affichage
            scores_plot = [s if s is not None else 0 for s in scores]
            fig.add_trace(go.Scatterpolar(
                r=scores_plot,
                theta=[f"Axe {a}" for a in AXES],
                fill="toself", name=row["nom_repondant"], opacity=0.5,
            ))
        fig.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 4])),
            showlegend=True, height=460,
        )
        st.plotly_chart(fig, use_container_width=True)

    # Detail par critere : distribution des statuts par question
    st.markdown("---")
    st.subheader("🔍 Detail par critere : distribution des statuts par question")
    selected_c = st.selectbox(
        "Choisir un critere",
        options=[c[0] for c in CRITERES],
        format_func=lambda x: f"{x} - {dict((c[0], c[2]) for c in CRITERES)[x]}",
    )
    a = agg[selected_c]
    if a["nb_reponses"] == 0:
        st.info("Aucune reponse sur ce critere pour l'instant.")
    else:
        st.markdown(
            f"**Score moyen** : `{a['score_moyen']:.2f} / 4`  ·  "
            f"min `{a['score_min']:.2f}` / max `{a['score_max']:.2f}`  ·  "
            f"**{a['nb_reponses']} reponses**  ·  "
            f"Niveau : {qualif(a['score_moyen'])}"
        )
        for q in CHECKLISTS[selected_c]:
            qa = a["questions"][q["id"]]
            total = qa["obs"] + qa["part"] + qa["non"]
            q_score = ((qa["obs"] * 1.0 + qa["part"] * 0.5) / total * 4) if total > 0 else None
            with st.container():
                st.markdown(f"**{q['id']}** {q['label']}")
                if total > 0:
                    cols = st.columns(4)
                    cols[0].metric("✅ Observe", qa["obs"])
                    cols[1].metric("🟡 Partiel", qa["part"])
                    cols[2].metric("❌ Non observe", qa["non"])
                    cols[3].metric("Score q./4",
                                   f"{q_score:.2f}" if q_score is not None else "—")
                    if qa["preuves"]:
                        with st.expander(f"📎 {len(qa['preuves'])} preuves collectees"):
                            for preuve in qa["preuves"]:
                                st.markdown(f"- {preuve}")
                else:
                    st.caption("Aucune reponse a cette question.")
                st.markdown("---")

    # Export Word
    st.markdown("---")
    st.subheader("📥 Exporter la synthese")
    if org_filtree:
        st.caption(
            f"L'export ne contient que les reponses pour **{org_filtree}** "
            f"({len(df)} repondant{'s' if len(df) > 1 else ''})."
        )
        org_slug = "".join(c if c.isalnum() else "_" for c in org_filtree)[:40]
        filename = f"Synthese_RTE_{org_slug}_{date.today().isoformat()}.docx"
    else:
        st.caption(
            f"L'export contient toutes les reponses de toutes les organisations "
            f"({len(df_all)} reponse{'s' if len(df_all) > 1 else ''})."
        )
        filename = f"Synthese_RTE_{date.today().isoformat()}.docx"

    buf = export_synthese_docx(df)
    st.download_button(
        label="📄 Telecharger la synthese en Word (.docx)",
        data=buf,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )


# ============================================================
# UI : PAGE ADMINISTRATION
# ============================================================

def page_admin():
    st.title("⚙️ Administration")
    df = load_responses()

    st.subheader(f"Reponses enregistrees : {len(df)}")
    if len(df) == 0:
        st.info("Aucune reponse enregistree.")
        return

    st.warning(
        "⚠️ **Pense a sauvegarder regulierement** les reponses. "
        "Si l'app est redeployee, la base est reinitialisee. "
        "Telecharge un backup CSV ci-dessous au moins une fois par semaine."
    )
    col1, col2 = st.columns(2)
    with col1:
        csv_data = df.to_csv(index=False).encode("utf-8")
        st.download_button(
            "💾 Telecharger backup CSV", data=csv_data,
            file_name=f"backup_RTE_{date.today().isoformat()}.csv",
            mime="text/csv", type="primary", use_container_width=True,
        )
    with col2:
        json_data = df.to_json(orient="records", force_ascii=False, indent=2).encode("utf-8")
        st.download_button(
            "📦 Backup JSON complet", data=json_data,
            file_name=f"backup_RTE_{date.today().isoformat()}.json",
            mime="application/json", use_container_width=True,
        )

    st.markdown("---")
    display = df[["id", "date_soumission", "nom_repondant", "fonction",
                  "organisation", "email"]].copy()
    st.dataframe(display, hide_index=True, use_container_width=True)

    st.markdown("---")
    st.subheader("Supprimer une reponse")
    rid = st.number_input("ID a supprimer", min_value=1, step=1)
    if st.button("🗑️ Supprimer cette reponse"):
        delete_response(rid)
        st.success(f"Reponse #{rid} supprimee. Rechargez la page.")

    st.markdown("---")
    st.subheader("⚠️ Reinitialiser toute la base")
    st.warning("Cette action supprime TOUTES les reponses. Irreversible.")
    confirm = st.text_input("Tapez 'SUPPRIMER' pour confirmer")
    if st.button("🔥 Reinitialiser la base"):
        if confirm == "SUPPRIMER":
            reset_db()
            st.success("Base reinitialisee.")
        else:
            st.error("Tapez 'SUPPRIMER' exactement pour confirmer.")


# ============================================================
# MAIN
# ============================================================

def hide_sidebar_css():
    """Cache la sidebar et son bouton de deroulement (mode public/collegues)."""
    st.markdown("""
    <style>
        section[data-testid="stSidebar"] { display: none !important; }
        div[data-testid="collapsedControl"] { display: none !important; }
        .block-container { padding-left: 3rem !important; padding-right: 3rem !important; }
    </style>
    """, unsafe_allow_html=True)


def page_admin_share():
    """Affiche les liens a partager : URL publique vs URL admin."""
    st.markdown("### 🔗 Liens a partager")
    # Construction de l'URL courante
    try:
        url_admin_param = st.query_params.get("admin", "")
    except Exception:
        url_admin_param = ""

    st.info(
        "**👥 Lien pour vos collegues** (a envoyer par mail) :  \n"
        "Copiez l'URL actuelle de votre navigateur SANS le bout `?admin=...`. "
        "Ils verront uniquement le formulaire (pas le dashboard, pas l'admin)."
    )
    st.success(
        f"**🔑 Votre lien admin** (a garder pour vous) :  \n"
        f"Cette URL avec `?admin={url_admin_param}` a la fin vous donne acces "
        f"au Dashboard et a l'Admin. Mettez-la en favori."
    )
    st.markdown("---")
    st.markdown(
        "**💡 Pour changer le code secret :**  \n"
        "Modifiez la valeur de `DEFAULT_ADMIN_CODE` dans `app_rte.py` (ligne ~40), "
        "OU sur Streamlit Cloud : Settings → Secrets → ajoutez :  \n"
        "```\nadmin_code = \"mon_nouveau_code\"\n```"
    )


def main():
    st.set_page_config(
        page_title=APP_TITLE, page_icon=APP_ICON, layout="wide",
        initial_sidebar_state="expanded",
    )
    apply_custom_css()
    init_db()

    # ============================================================
    # CONTROLE D'ACCES via parametre URL ?admin=<code>
    # ============================================================
    try:
        provided_code = st.query_params.get("admin", "")
    except Exception:
        provided_code = ""
    is_admin = (provided_code == get_admin_code())

    if not is_admin:
        # === MODE PUBLIC / COLLEGUES : formulaire seul, sidebar cachee ===
        hide_sidebar_css()
        page_formulaire()
        return

    # === MODE ADMIN : sidebar complete avec toutes les pages ===
    with st.sidebar:
        st.title("🧭 Navigation")
        st.markdown("**Grille de maturite RTE**")
        st.caption("Responsabilite Territoriale des Entreprises")
        st.markdown(
            "<div style='background:rgba(225,182,96,0.15);padding:8px 12px;"
            "border-radius:6px;border-left:3px solid #E1B660;margin:8px 0;'>"
            "<b style='color:#E1B660'>🔑 MODE ADMIN</b><br>"
            "<span style='font-size:11px;opacity:0.85'>Vos collegues n'ont acces "
            "qu'au Formulaire.</span>"
            "</div>",
            unsafe_allow_html=True
        )
        st.markdown("---")
        page = st.radio(
            "Choisir la page",
            options=["📝 Formulaire", "📊 Dashboard", "⚙️ Admin", "🔗 Partage"],
            label_visibility="collapsed",
        )
        st.markdown("---")
        df = load_responses()
        st.metric("Reponses recues", len(df))
        st.markdown("---")
        st.markdown(
            "**Notation auto** :  \n"
            "✅ Observe = 1 pt  \n"
            "🟡 Partiel = 0,5 pt  \n"
            "❌ Non observe = 0 pt  \n  \n"
            "*Score critere = (moyenne / nb questions) × 4*"
        )
        st.markdown("---")
        st.caption("Stage M2 IE IAE Poitiers")
        st.caption("Chaire TerrESS - Bordeaux Sciences Agro")

    if page == "📝 Formulaire":
        page_formulaire()
    elif page == "📊 Dashboard":
        page_dashboard()
    elif page == "⚙️ Admin":
        page_admin()
    elif page == "🔗 Partage":
        page_admin_share()


if __name__ == "__main__":
    main()
